def get_features_from(s):
    '''
    给一段string，选出其中的features来
    '''
    features_str = s.replace("'", '').split('features=')[1].split(' state')[0]
    return features_str.split(';')

def get_text_part(f_name):
    '''
    将xml文件中的内容读出来，然后保存成list格式
    '''
    with open(f_name, 'r', encoding = 'utf8') as f:
        data = f.read()
        data = data.split('<body>')[1].split('</body>')[0].strip()
        data = data.replace('&apos;', "'")
        data = data.replace('&amp;', '&')
        all_datas = []
        for xx in data.split('<')[1:]:
            if '/segment>' in xx:
                this_text = xx.replace('/segment>', '')
                features = ''
            else :
                this_text = xx.split('>')[1]
                features = get_features_from(xx)
            
            if  not this_text.replace(' ', '') == 'I' and (this_text.replace(' ', '') == '' or len(this_text.replace(' ', '')) == 1):
                    features = ['something_else']
                
            data_point = []
            data_point = [this_text, features]
            all_datas.append(data_point) 
            
    return all_datas 

from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.enum.text import WD_BREAK
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Cm
import traceback
from docx.oxml.shared import OxmlElement, qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

def set_cell_vertical_alignment(cell, align="center"): 
    try:   
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcValign = OxmlElement('w:vAlign')  
        tcValign.set(qn('w:val'), align)  
        tcPr.append(tcValign)
        return True 
    except:
        traceback.print_exc()             
        return False

def save_data_to_file(write_out, f_name):
    doc = Document()

    # 表格内文本格式
    obj_styles = doc.styles
    obj_charstyle2 = obj_styles.add_style('table_style', WD_STYLE_TYPE.CHARACTER)
    obj_font2 = obj_charstyle2.font
    obj_font2.size = Pt(9)
    obj_font2.name = u'Times New Roman'
    r2 = doc.styles['table_style']._element
    r2.rPr.rFonts.set(qn('w:eastAsia'), u'楷体')

    table = doc.add_table(rows = 1, cols = 4)
    table.style = 'TableGrid'
    hdr_cells = table.rows[0].cells
    hdr_content = ['Independent Clause', 'textual', 'interpersonal', 'Theme Status']

    cell_width = [13, 1, 1, 3 ]
    for ii, h_c in enumerate(hdr_content):
        hdr_cells[ii].text = ''
        hdr_cells[ii].paragraphs[0].add_run(h_c, style = 'table_style')
        hdr_cells[ii].width = Cm(cell_width[ii])
        hdr_cells[ii].height = Cm(0.8)


    for each_c in hdr_cells:
        each_c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_vertical_alignment(each_c) 

    has_textual = False
    has_interpersonal = False
    add_new_row = True
    for xx in write_out:
        if add_new_row:
            current_row = table.add_row()

        if not xx[1] == '':
            add_new_row = False

            row_cells = current_row.cells
            # row_cells[0].text = ''
            this_run = row_cells[0].paragraphs[0].add_run(xx[0], style = 'table_style' )
            this_run.underline = True

            if 'textual' in xx[1]: 
                has_textual = True    
                this_run.bold = True
            elif 'interpersonal' in xx[1]:
                has_interpersonal = True
                this_run.italic = True  
            elif not 'something_else' in xx[1]:
                experiential_type = xx[1]
        else:
            add_new_row = True
            another_run = row_cells[0].paragraphs[0].add_run(xx[0], style = 'table_style' )

            next_content = {
                'has_textual' : '-',
                'has_interpersonal' : '-',
                'status' : '-'.join(experiential_type[1:])
            }

            if has_textual : next_content['has_textual'] = '+'
            if has_interpersonal : next_content['has_interpersonal'] = '+'

            for ii, the_content in enumerate(['has_textual', 'has_interpersonal','status' ]):
                row_cells[1 + ii ].text = ''
                row_cells[1 + ii].paragraphs[0].add_run(next_content[the_content], style = 'table_style' )

            for jj in range(4):
                row_cells[jj].width = Cm(cell_width[ii])

            has_textual = False
            has_interpersonal = False
    for each_r in table.rows:
        tr = each_r._tr
        trPr = tr.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), "512")
        trHeight.set(qn('w:hRule'), "atLeast")
        trPr.append(trHeight)

        #cell = each_r.cells[0]
        for cell in each_r.cells:
            set_cell_vertical_alignment(cell)
    doc.save(f_name)    